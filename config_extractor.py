"""
Context-Aware Configuration Extraction Agent

Features:
- Hierarchical multi-pass processing
- Long-term memory store for context preservation
- Document structure understanding
- Cross-reference resolution
- Progressive refinement of configurations
"""

import asyncio
import base64
import hashlib
import io
import json
import logging
import os
import re
import sqlite3
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional

import openai
import yaml
# Third-party imports - Document Processing
from PIL import Image
from docx import Document, ImagePart
from docx.shape import InlineShape
from langchain_core.prompts import ChatPromptTemplate

from pypdf import PdfReader

try:
    from fitz import fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        import PyMuPDF as fitz

# Third-party imports - OpenAI and LangChain
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Third-party imports - Data Processing

# Third-party imports - CLI and Interface
from rich.console import Console
from rich.progress import (
    Progress
)
from rich.logging import RichHandler

# Third-party imports - Configuration and Validation
from pydantic import BaseModel, Field

# Local imports
import config

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(message)s",
    datefmt="[%X]",
    handlers=[RichHandler(rich_tracebacks=True)]
)

# Logger instance
logger = logging.getLogger("config_extractor")

# Constants
SUPPORTED_EXTENSIONS = {
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'docx',
    '.txt': 'text',
    '.json': 'json',
    '.yaml': 'yaml',
    '.yml': 'yaml',
    '.mp4': 'video',
    '.avi': 'video',
    '.mov': 'video',
    '.mkv': 'video',
    '.mp3': 'audio',
    '.wav': 'audio',
    '.m4a': 'audio'
}


# Previous schema definitions remain the same...

class MemoryStore:
    """Simplified memory store without vector operations."""

    def __init__(self, db_path: str = "config_memory.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Initialize SQLite database for storing metadata and relationships."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS config_history (
                    doc_hash TEXT PRIMARY KEY,
                    doc_path TEXT,
                    last_processed TIMESTAMP,
                    config_json TEXT,
                    metadata_json TEXT
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS config_relationships (
                    source_hash TEXT,
                    target_hash TEXT,
                    relationship_type TEXT,
                    confidence FLOAT,
                    FOREIGN KEY (source_hash) REFERENCES config_history(doc_hash),
                    FOREIGN KEY (target_hash) REFERENCES config_history(doc_hash)
                )
            """)

    def store_context(self, doc_path: Path, text: str, config: Dict[str, Any], metadata: Dict[str, Any]):
        """Store document context and configuration with relationships."""
        doc_hash = self._compute_hash(text)

        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                "INSERT OR REPLACE INTO config_history VALUES (?, ?, ?, ?, ?)",
                (
                    doc_hash,
                    str(doc_path),
                    datetime.now().isoformat(),
                    json.dumps(config),
                    json.dumps(metadata)
                )
            )

    def find_related_contexts(self, text: str, k: int = 3) -> List[Dict[str, Any]]:
        """Find related configurations based on simple matching."""
        doc_hash = self._compute_hash(text)
        related_contexts = []

        with sqlite3.connect(self.db_path) as conn:
            rows = conn.execute(
                """
                SELECT config_json, metadata_json 
                FROM config_history 
                WHERE doc_hash != ? 
                ORDER BY last_processed DESC 
                LIMIT ?
                """,
                (doc_hash, k)
            ).fetchall()

            for row in rows:
                related_contexts.append({
                    "config": json.loads(row[0]),
                    "metadata": json.loads(row[1]),
                    "similarity_score": 1.0  # Placeholder score
                })

        return related_contexts

    @staticmethod
    def _compute_hash(text: str) -> str:
        """Compute document hash for identification."""
        return hashlib.sha256(text.encode()).hexdigest()


class DocumentStructureAnalyzer:
    """Analyzes document structure and relationships between sections."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ";", ",", " "]
        )

    async def analyze_structure(self, text: str) -> Dict[str, Any]:
        """Perform structural analysis of the document."""
        # Split document into sections
        sections = self._identify_sections(text)

        # Analyze relationships between sections
        relationships = self._analyze_relationships(sections)

        # Identify configuration-relevant sections
        config_sections = self._identify_config_sections(sections)

        return {
            "sections": sections,
            "relationships": relationships,
            "config_sections": config_sections
        }

    def _identify_sections(self, text: str) -> List[Dict[str, Any]]:
        """Identify logical sections in the document."""
        sections = []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Process each chunk as a section
        for i, chunk in enumerate(chunks):
            # Try to identify section type and heading
            section_type = self._determine_section_type(chunk)
            heading = self._extract_heading(chunk)
            
            section = {
                "id": i,
                "type": section_type,
                "heading": heading,
                "content": chunk,
                "length": len(chunk)
            }
            sections.append(section)
        
        return sections

    def _determine_section_type(self, text: str) -> str:
        """Determine the type of section based on content analysis."""
        # Simple heuristic-based section type determination
        text_lower = text.lower()
        
        if any(word in text_lower for word in ["config", "configuration", "settings", "parameters"]):
            return "configuration"
        elif any(word in text_lower for word in ["database", "db", "sql", "nosql"]):
            return "database"
        elif any(word in text_lower for word in ["api", "endpoint", "rest", "service"]):
            return "api"
        elif any(word in text_lower for word in ["security", "auth", "authentication", "permission"]):
            return "security"
        else:
            return "general"

    def _extract_heading(self, text: str) -> Optional[str]:
        """Extract section heading if present."""
        # Simple heading extraction logic
        lines = text.split('\n')
        if lines:
            # Look for potential heading in first few lines
            for line in lines[:3]:
                line = line.strip()
                # Check if line looks like a heading
                if line and len(line) < 100 and line.strip().endswith(':'):
                    return line.strip(':')
        return None

    def _analyze_relationships(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Analyze relationships between sections."""
        relationships = []
        
        for i, section in enumerate(sections):
            # Look for references to other sections
            for j, other_section in enumerate(sections):
                if i != j:
                    if self._sections_are_related(section, other_section):
                        relationship = {
                            "source": section["id"],
                            "target": other_section["id"],
                            "type": "reference",
                            "confidence": 0.8  # Placeholder confidence score
                        }
                        relationships.append(relationship)
        
        return relationships

    def _sections_are_related(self, section1: Dict[str, Any], section2: Dict[str, Any]) -> bool:
        """Determine if two sections are related."""
        # Simple relationship detection based on content overlap
        content1 = section1["content"].lower()
        content2 = section2["content"].lower()
        
        # Check for common configuration terms
        config_terms = ["config", "setting", "parameter", "option"]
        
        # Check if both sections contain similar configuration terms
        return any(term in content1 and term in content2 for term in config_terms)

    def _identify_config_sections(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Identify sections that contain configuration information."""
        config_sections = []
        
        for section in sections:
            if section["type"] == "configuration":
                config_sections.append(section)
            elif self._contains_config_patterns(section["content"]):
                config_sections.append(section)
        
        return config_sections

    def _contains_config_patterns(self, text: str) -> bool:
        """Check if text contains configuration patterns."""
        # Common configuration patterns
        patterns = [
            r'(?i)config[uration]*\s*=',
            r'(?i)setting[s]*\s*=',
            r'(?i)param[s|eter]*\s*=',
            r'(?i)options?\s*=',
            r'(?i)env[ironment]*\s*=',
            r'(?i)var[iable]*\s*='
        ]
        
        return any(re.search(pattern, text) for pattern in patterns)


class DatabaseConfig(BaseModel):
    """Database configuration schema."""
    host: Optional[str] = Field(default=None, description="Database host address")
    port: Optional[int] = Field(default=None, description="Database port number", ge=1, le=65535)
    username: Optional[str] = Field(default=None, description="Database username")
    password: Optional[str] = Field(default=None, description="Database password")

    class Config:
        extra = "forbid"


class ServerConfig(BaseModel):
    """Server configuration schema."""
    ip: Optional[str] = Field(default=None, description="Server IP address")
    uptime: Optional[str] = Field(default=None, description="Server uptime requirement")

    class Config:
        extra = "forbid"


class ConfigSchema(BaseModel):
    """Main configuration schema with nested models."""
    database: Optional[DatabaseConfig] = Field(default=None, description="Database configuration settings")
    server: Optional[ServerConfig] = Field(default=None, description="Server configuration settings")

    class Config:
        extra = "forbid"


class SchemaValidator:
    """Custom JSON schema validator implementation."""

    def __init__(self, schema: dict):
        self.schema = schema

    def validate(self, instance: Any) -> tuple[bool, list[str]]:
        """
        Validate an instance against the schema.

        Args:
            instance: The data to validate

        Returns:
            tuple[bool, list[str]]: (is_valid, list of validation errors)
        """
        self.errors = []
        is_valid = self._validate_against_schema(instance, self.schema, path="$")
        return is_valid, self.errors

    def _validate_against_schema(self, instance: Any, schema: dict, path: str) -> bool:
        """
        Recursively validate instance against schema.

        Args:
            instance: The data to validate
            schema: The schema to validate against
            path: Current JSON path for error reporting

        Returns:
            bool: True if valid, False otherwise
        """
        if not isinstance(schema, dict):
            self.errors.append(f"Invalid schema at {path}: schema must be an object")
            return False

        # Check type
        schema_type = schema.get("type")
        if not schema_type:
            return True  # No type constraint

        # Validate type
        if not self._validate_type(instance, schema_type, path):
            return False

        # Validate based on type
        if schema_type == "object":
            return self._validate_object(instance, schema, path)
        elif schema_type == "array":
            return self._validate_array(instance, schema, path)
        elif schema_type in ["string", "number", "integer", "boolean"]:
            return True  # Basic type validation already done
        else:
            self.errors.append(f"Unsupported type '{schema_type}' at {path}")
            return False

    def _validate_type(self, instance: Any, expected_type: str, path: str) -> bool:
        """Validate type of instance."""
        if expected_type == "object":
            if not isinstance(instance, dict):
                self.errors.append(f"Expected object at {path}, got {type(instance).__name__}")
                return False
        elif expected_type == "array":
            if not isinstance(instance, list):
                self.errors.append(f"Expected array at {path}, got {type(instance).__name__}")
                return False
        elif expected_type == "string":
            if not isinstance(instance, str):
                self.errors.append(f"Expected string at {path}, got {type(instance).__name__}")
                return False
        elif expected_type == "number":
            if not isinstance(instance, (int, float)):
                self.errors.append(f"Expected number at {path}, got {type(instance).__name__}")
                return False
        elif expected_type == "integer":
            if not isinstance(instance, int):
                self.errors.append(f"Expected integer at {path}, got {type(instance).__name__}")
                return False
        elif expected_type == "boolean":
            if not isinstance(instance, bool):
                self.errors.append(f"Expected boolean at {path}, got {type(instance).__name__}")
                return False
        return True

    def _validate_object(self, instance: dict, schema: dict, path: str) -> bool:
        """Validate object against schema."""
        is_valid = True

        # Check required properties
        required = schema.get("required", [])
        for prop in required:
            if prop not in instance:
                self.errors.append(f"Missing required property '{prop}' at {path}")
                is_valid = False

        # Validate properties
        properties = schema.get("properties", {})
        for prop, value in instance.items():
            if prop in properties:
                prop_path = f"{path}.{prop}"
                if not self._validate_against_schema(value, properties[prop], prop_path):
                    is_valid = False

        return is_valid

    def _validate_array(self, instance: list, schema: dict, path: str) -> bool:
        """Validate array against schema."""
        is_valid = True

        # Validate items
        if "items" in schema:
            for i, item in enumerate(instance):
                item_path = f"{path}[{i}]"
                if not self._validate_against_schema(item, schema["items"], item_path):
                    is_valid = False

        return is_valid


class BaseConfigExtractor:
    """Base class for configuration extraction."""

    def __init__(self, model_name: str = "gpt-4", schema_path: Optional[str] = None):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.llm = ChatOpenAI(model_name=model_name, temperature=0)
        self.schema = self._load_schema(schema_path) if schema_path else {}
        self.validator = SchemaValidator(self.schema) if self.schema else None
        self.setup_prompts()

    def _load_schema(self, schema_path: str) -> Dict[str, Any]:
        """Load JSON schema from file."""
        try:
            with open(schema_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            self.logger.error(f"Error loading schema: {e}")
            raise ValueError(f"Invalid schema file: {e}")

    def setup_prompts(self):
        """Setup base prompts. Override in subclasses."""
        pass

    async def process_document(self, file_path: Path, doc_type: str) -> Dict[str, Any]:
        """Base document processing method."""
        try:
            # Read document content
            text = await self._read_document(file_path, doc_type)

            # First pass analysis
            first_pass_results = await self._first_pass_analysis(text)

            # Structure analysis
            structure_analysis = await self._analyze_structure(text)

            # Find related contexts
            related_contexts = await self._find_related_contexts(text)

            # Second pass extraction
            config_data = await self._second_pass_extraction(
                text,
                first_pass_results,
                structure_analysis,
                related_contexts
            )

            return config_data

        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {e}")
            return self._create_empty_config()

    async def _read_document(self, file_path: Path, doc_type: str) -> str:
        """Read document content based on type."""
        # Implementation would vary by document type
        raise NotImplementedError

    async def _first_pass_analysis(self, text: str) -> Dict[str, Any]:
        """Perform first pass analysis of document."""
        raise NotImplementedError

    async def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure."""
        raise NotImplementedError

    async def _find_related_contexts(self, text: str) -> List[Dict[str, Any]]:
        """Find related configuration contexts."""
        raise NotImplementedError

    async def _second_pass_extraction(
            self,
            text: str,
            first_pass_results: Dict[str, Any],
            structure_analysis: Dict[str, Any],
            related_contexts: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Extract configurations in second pass."""
        raise NotImplementedError

    def _create_empty_config(self) -> Dict[str, Any]:
        """Create an empty configuration matching the schema structure."""
        if not self.schema:
            return {}

        def create_empty_value(schema_part: Dict[str, Any]) -> Any:
            schema_type = schema_part.get("type")
            if schema_type == "object":
                properties = schema_part.get("properties", {})
                required = schema_part.get("required", [])
                obj = {}
                for prop, prop_schema in properties.items():
                    if prop in required:
                        obj[prop] = create_empty_value(prop_schema)
                return obj
            elif schema_type == "array":
                return []
            elif schema_type == "string":
                return ""
            elif schema_type == "boolean":
                return False
            elif schema_type in ["number", "integer"]:
                return 0
            else:
                return None

        return create_empty_value(self.schema)


class ConfigExtractor(BaseConfigExtractor):
    """Configuration extractor with video and image support."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.memory_store = MemoryStore()
        self.structure_analyzer = DocumentStructureAnalyzer()
        # OpenAI client for transcription
        self.openai_client = openai.OpenAI()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ";", ",", " "]
        )
        from langchain.prompts import ChatPromptTemplate
        self.setup_prompts()

    # Replace the existing _read_docx method with the new implementation
    async def _read_docx(self, file_path: Path) -> str:
        """Read content from a DOCX/DOC file with format detection and conversion."""
        try:
            # First try to detect if it's really a DOC file despite the extension
            is_doc = self._is_doc_format(file_path)

            if is_doc:
                logger.info(f"Detected DOC format for {file_path}")
                self.logger.info(f"Detected DOC format for {file_path}")
                # For DOC files, try conversion methods
                return await self._handle_doc_file(file_path)
            else:
                logger.info(f"Detected DOCX format for {file_path}")
                # For DOCX files, use python-docx directly
                return await self._handle_docx_file(file_path)

        except Exception as e:
            self.logger.error(f"Error reading document file {file_path}: {e}")
            raise

    async def _handle_docx_file(self, file_path: Path) -> str:
        """Handle DOCX format files using python-docx."""
        try:
            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)
        
        
        

        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise

    async def _handle_doc_file(self, file_path: Path) -> str:
        """Handle DOC format files using available conversion methods."""
        try:
            # First try LibreOffice conversion
            converted_path = await self._convert_doc_to_docx(file_path)
            if converted_path:
                try:
                    text = await self._handle_docx_file(converted_path)
                    os.remove(converted_path)
                    return text
                except:
                    os.remove(converted_path)
                    raise

            # If conversion fails, try antiword
            if self._has_antiword():
                self.logger.info("Attempting to read with antiword")
                return await self._read_with_antiword(file_path)

            # Finally try catdoc
            if self._has_catdoc():
                self.logger.info("Attempting to read with catdoc")
                return await self._read_with_catdoc(file_path)

            raise Exception("No available method to read DOC file. Please install LibreOffice, antiword, or catdoc.")

        except Exception as e:
            raise Exception(f"Failed to read DOC file: {str(e)}")

    # Add all the helper methods for DOC/DOCX handling
    def _is_doc_format(self, file_path: Path) -> bool:
        """
        Check if the file is in old DOC format by looking at its magic numbers and structure.

        Both DOC and DOCX files use the same compound file binary format (CFBF) magic number,
        but have different internal structures. We need to check both.
        """
        # Define magic numbers
        CFBF_MAGIC = b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'  # Compound File Binary Format
        ZIP_MAGIC = b'PK\x03\x04'  # ZIP format (used by DOCX)

        try:
            with open(file_path, 'rb') as f:
                magic = f.read(8)

                # If it's a ZIP file (DOCX), it's not a DOC
                if magic.startswith(ZIP_MAGIC):
                    return False

                # If it doesn't have the CFBF magic number, it's not a DOC
                if magic != CFBF_MAGIC:
                    return False

                # Additional check: DOC files have specific patterns
                # Move to the start of the file
                f.seek(0)
                content = f.read(2048)  # Read first 2KB for analysis

                # DOCX files (even with CFBF magic) won't have these patterns
                doc_patterns = [
                    b'Microsoft Word',
                    b'MSWordDoc',
                    b'Word.Document.'
                ]

                return any(pattern in content for pattern in doc_patterns)

        except Exception as e:
            self.logger.warning(f"Error checking file format: {e}")
            # If we can't check the binary content, use extension as fallback
            # but log a warning
            is_doc = file_path.suffix.lower() == '.doc'
            if is_doc:
                self.logger.warning("Falling back to extension-based detection")
            return is_doc

    async def _convert_doc_to_docx(self, file_path: Path) -> Optional[Path]:
        """Convert DOC to DOCX using available tools."""
        # Try LibreOffice first
        if self._has_libreoffice():
            try:
                temp_dir = tempfile.mkdtemp()
                output_path = Path(temp_dir) / f"{file_path.stem}.docx"

                # Use LibreOffice to convert
                process = await asyncio.create_subprocess_exec(
                    'soffice',
                    '--headless',
                    '--convert-to',
                    'docx',
                    str(file_path),
                    '--outdir',
                    temp_dir,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                await process.communicate()

                if output_path.exists():
                    return output_path
            except:
                self.logger.warning("LibreOffice conversion failed, trying alternatives")

        return None

    def _has_libreoffice(self) -> bool:
        """Check if LibreOffice is available."""
        try:
            subprocess.run(['soffice', '--version'],
                           stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE)
            return True
        except:
            return False

    def _has_antiword(self) -> bool:
        """Check if antiword is available."""
        try:
            subprocess.run(['antiword', '-v'],
                           stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE)
            return True
        except:
            return False

    def _has_catdoc(self) -> bool:
        """Check if catdoc is available."""
        try:
            subprocess.run(['catdoc', '-v'],
                           stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE)
            return True
        except:
            return False

    async def _read_with_antiword(self, file_path: Path) -> str:
        """Read DOC file using antiword."""
        try:
            process = await asyncio.create_subprocess_exec(
                'antiword',
                str(file_path),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await process.communicate()
            if process.returncode == 0:
                return stdout.decode('utf-8', errors='ignore')
            else:
                raise Exception(f"antiword failed: {stderr.decode()}")
        except Exception as e:
            raise Exception(f"Error using antiword: {str(e)}")

    async def _read_with_catdoc(self, file_path: Path) -> str:
        """Read DOC file using catdoc."""
        try:
            process = await asyncio.create_subprocess_exec(
                'catdoc',
                str(file_path),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            stdout, stderr = await process.communicate()
            if process.returncode == 0:
                return stdout.decode('utf-8', errors='ignore')
            else:
                raise Exception(f"catdoc failed: {stderr.decode()}")
        except Exception as e:
            raise Exception(f"Error using catdoc: {str(e)}")

    async def _read_document(self, file_path: Path, doc_type: str) -> str:
        """Read document content based on type."""
        try:
            if doc_type == 'pdf':
                return await self._read_pdf(file_path)
            elif doc_type == 'docx':
                return await self._read_docx(file_path)
            elif doc_type in ['video', 'audio']:
                return await self._process_media_file(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
        except Exception as e:
            self.logger.error(f"Error reading document {file_path}: {e}")
            raise

    async def _process_media_file(self, file_path: Path) -> str:
        """Process video/audio file using Whisper API for transcription."""
        try:
            with open(file_path, "rb") as media_file:
                transcription = await self.openai_client.audio.transcriptions.create(
                    file=media_file,
                    model="whisper-1",
                    response_format="text",
                    language="en"
                )
                return transcription
        except Exception as e:
            self.logger.error(f"Error transcribing media file {file_path}: {e}")
            raise

    async def _second_pass_extraction(
            self,
            text: str,
            first_pass_results: Dict[str, Any],
            structure_analysis: Dict[str, Any],
            related_contexts: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Extract configurations in second pass with chunking for long texts."""
        max_retries = 3
        retry_delay = 1  # seconds
        
        for attempt in range(max_retries):
            try:
            # Split text into manageable chunks if it's too long
                chunks = self.text_splitter.split_text(text)

                # Process each chunk
                configs = []
                for chunk in chunks:
                    messages = self.second_pass_prompt.format_messages(
                        text=chunk,
                        structure_summary=first_pass_results["document_summary"],
                        related_configs=json.dumps(related_contexts),
                        relationships=json.dumps(structure_analysis["relationships"]),
                        schema=json.dumps(self.schema, indent=2)
                    )

                    response = await self.llm.agenerate([messages])
                    try:
                        chunk_config = json.loads(response.generations[0][0].text)
                        configs.append(chunk_config)
                    except json.JSONDecodeError as e:
                        self.logger.warning(f"Failed to parse chunk configuration: {e}")
                        continue

                # Merge configurations from all chunks
                merged_config = self._merge_chunk_configs(configs)

                # Validate final configuration
                if self.validator:
                    is_valid, errors = self.validator.validate(merged_config)
                    if not is_valid:
                        self.logger.error("Validation errors in extracted configuration:")
                        for error in errors:
                            self.logger.error(f"  - {error}")
                        return self._create_empty_config()

                return merged_config

            except Exception as e:
                if attempt < max_retries - 1:
                    self.logger.warning(f"Attempt {attempt + 1} failed: {e}. Retrying...")
            return self._create_empty_config()

    def _merge_chunk_configs(self, configs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge configurations from multiple chunks with conflict resolution."""
        if not configs:
            return self._create_empty_config()

        merged = configs[0].copy()
        for config in configs[1:]:
            for key, value in config.items():
                if key not in merged:
                    merged[key] = value
                elif isinstance(value, dict) and isinstance(merged[key], dict):
                    merged[key] = self._merge_chunk_configs([merged[key], value])
                else:
                    # Use the most specific/detailed value
                    merged[key] = self._resolve_config_conflict(merged[key], value)

        return merged

    def _resolve_config_conflict(self, value1: Any, value2: Any) -> Any:
        """Resolve conflicts between two configuration values."""
        if value1 is None:
            return value2
        if value2 is None:
            return value1

        # Prefer non-default values
        default_values = {'localhost', '127.0.0.1', 'user', 'password', '0', '', '[]', '{}'}
        if str(value1).lower() in default_values and str(value2).lower() not in default_values:
            return value2
        if str(value2).lower() in default_values and str(value1).lower() not in default_values:
            return value1

        # Prefer more specific/detailed values
        if isinstance(value1, str) and isinstance(value2, str):
            if len(value2) > len(value1):
                return value2

        # Default to first value
        return value1

    async def _read_pdf(self, file_path: Path) -> str:
        """Read content from a PDF file using PyMuPDF (fitz)."""
        try:
            text_parts = []
            with fitz.open(file_path) as pdf_doc:
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    text_parts.append(page.get_text())

            return "\n".join(text_parts)
        except Exception as e:
            self.logger.error(f"Error reading PDF file {file_path}: {e}")
            raise

    async def _first_pass_analysis(self, text: str) -> Dict[str, Any]:
        """Perform first pass analysis of document."""
        try:
            # Extract potential configurations using regex and heuristics
            potential_configs = self._extract_potential_configs(text)

            # Create document summary
            document_summary = {
                "length": len(text),
                "potential_config_count": len(potential_configs),
                "document_summary": "Document contains configuration settings and parameters",
                "config_patterns": [conf['type'] for conf in potential_configs]
            }

            return {
                "document_summary": document_summary,
                "potential_configs": potential_configs
            }
        except Exception as e:
            self.logger.error(f"Error in first pass analysis: {e}")
            return {"document_summary": {}, "potential_configs": []}

    async def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure."""
        try:
            # Use the DocumentStructureAnalyzer for structure analysis
            structure_analyzer = DocumentStructureAnalyzer()
            return await structure_analyzer.analyze_structure(text)
        except Exception as e:
            self.logger.error(f"Error in structure analysis: {e}")
            return {"sections": [], "relationships": [], "config_sections": []}

    async def _find_related_contexts(self, text: str) -> List[Dict[str, Any]]:
        """Find related configuration contexts."""
        try:
            # Use memory store to find related contexts
            return self.memory_store.find_related_contexts(text)
        except Exception as e:
            self.logger.error(f"Error finding related contexts: {e}")
            return []

    def setup_prompts(self):
        """Setup prompts for the extractor."""
        # Define the base prompt for second pass extraction
        self.second_pass_prompt = ChatPromptTemplate.from_messages([
            ("system", "You are a configuration extraction assistant. "
                       "Your task is to extract configuration values that match the provided schema. "
                       "Consider the document structure and related configurations when extracting values."),
            ("user", "Extract configurations from the following text, considering these inputs:\n"
                     "Text: {text}\n"
                     "Structure Summary: {structure_summary}\n"
                     "Related Configs: {related_configs}\n"
                     "Document Relationships: {relationships}\n"
                     "Schema: {schema}\n"
                     "Please provide the extracted configurations in valid JSON format matching the schema.")
        ])

    def _extract_potential_configs(self, text: str) -> List[Dict[str, Any]]:
        """Extract potential configuration patterns from text using regex and heuristics."""
        potential_configs = []

        # Common configuration patterns
        patterns = {
            'connection_string': r'(?i)(connection[_\s]string|conn[_\s]str)[_\s]*=\s*["\']([^"\']+)["\']',
            'host': r'(?i)(hostname|host)[_\s]*=\s*["\']?([^"\'\s,}]+)',
            'port': r'(?i)port[_\s]*=\s*["\']?(\d+)',
            'credentials': r'(?i)(username|user|password|pwd)[_\s]*=\s*["\']([^"\']+)["\']',
            'paths': r'(?i)(path|directory|dir|folder)[_\s]*=\s*["\']?([^"\'\s,}]+)',
            'timeout': r'(?i)timeout[_\s]*=\s*["\']?(\d+)',
            'url': r'(?i)(https?://[^\s,}"\']+)',
            'key_value': r'([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([^,}\s]+)'
        }

        # Extract matches for each pattern
        for pattern_name, pattern in patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                config_entry = {
                    'type': pattern_name,
                    'key': match.group(1) if len(match.groups()) > 1 else pattern_name,
                    'value': match.group(2) if len(match.groups()) > 1 else match.group(1),
                    'context': text[max(0, match.start() - 50):min(len(text), match.end() + 50)]
                }
                potential_configs.append(config_entry)

        # Look for JSON-like structures
        json_pattern = r'{[^{}]*}'
        json_matches = re.finditer(json_pattern, text)
        for match in json_matches:
            try:
                json_str = match.group(0)
                json_obj = json.loads(json_str)
                if isinstance(json_obj, dict):
                    potential_configs.append({
                        'type': 'json_structure',
                        'value': json_obj,
                        'context': text[max(0, match.start() - 50):min(len(text), match.end() + 50)]
                    })
            except json.JSONDecodeError:
                continue

        # Look for YAML-like structures
        yaml_pattern = r'(?m)^\s*[A-Za-z_][A-Za-z0-9_]*:\s*\n(?:\s+[^\n]+\n)+'
        yaml_matches = re.finditer(yaml_pattern, text)
        for match in yaml_matches:
            try:
                yaml_str = match.group(0)
                yaml_obj = yaml.safe_load(yaml_str)
                if isinstance(yaml_obj, dict):
                    potential_configs.append({
                        'type': 'yaml_structure',
                        'value': yaml_obj,
                        'context': text[max(0, match.start() - 50):min(len(text), match.end() + 50)]
                    })
            except yaml.YAMLError:
                continue

        # Filter out likely false positives
        filtered_configs = []
        for config in potential_configs:
            # Skip if value is a common word or placeholder
            if isinstance(config.get('value'), str):
                value = config['value'].lower()
                if value in {'true', 'false', 'none', 'null', 'undefined', 'example', 'test'}:
                    continue

            # Skip if value is too short and looks like a common word
            if isinstance(config.get('value'), str) and len(config['value']) < 3:
                continue

            filtered_configs.append(config)

        return filtered_configs


try:
    from fitz import fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        import PyMuPDF as fitz


# Usage
# Supported file extensions
class ImageExtractor:
    """Handles extraction of images from different document types."""

    def __init__(self):
        """Initialize the ImageExtractor with temporary directory and logging."""
        self.temp_dir = Path(tempfile.mkdtemp(prefix="config_images_"))
        self.image_hashes = set()  # Track unique images
        self.logger = logging.getLogger("ImageExtractor")

        # Ensure temp directory exists
        self.temp_dir.mkdir(parents=True, exist_ok=True)

        # Set minimum dimensions for valid images
        self.min_width = 50
        self.min_height = 50

    async def extract_images(self, file_path: Path, doc_type: str) -> List[Path]:
        """
        Extract images from document based on type.

        Args:
            file_path: Path to the document
            doc_type: Type of document ('pdf', 'docx', etc.)

        Returns:
            List of paths to extracted images
        """
        try:
            if doc_type == 'pdf':
                return await self._extract_pdf_images(file_path)
            elif doc_type == 'docx':
                return await self._extract_docx_images(file_path)
            else:
                self.logger.warning(f"Unsupported document type: {doc_type}")
                return []
        except Exception as e:
            self.logger.error(f"Error extracting images from {file_path}: {e}")
            return []

    async def _extract_pdf_images(self, pdf_path: Path) -> List[Path]:
        """
        Extract images from PDF file using pypdf.

        Args:
            pdf_path: Path to PDF file

        Returns:
            List of paths to extracted images
        """
        image_files = []
        try:
            reader = PdfReader(str(pdf_path))

            for page_num, page in enumerate(reader.pages):
                if "/Resources" in page and "/XObject" in page["/Resources"]:
                    xObject = page["/Resources"]["/XObject"].get_object()

                    for obj in xObject:
                        if xObject[obj]["/Subtype"] == "/Image":
                            try:
                                data = xObject[obj].get_data()
                                img_hash = hashlib.md5(data).hexdigest()

                                # Skip if we've seen this image before
                                if img_hash in self.image_hashes:
                                    continue

                                self.image_hashes.add(img_hash)

                                # Determine image format
                                format_mapping = {
                                    "/DCTDecode": "jpg",
                                    "/FlateDecode": "png",
                                    "/JPXDecode": "jp2",
                                    "/CCITTFaxDecode": "tiff"
                                }

                                filter_type = xObject[obj]["/Filter"]
                                if isinstance(filter_type, list):
                                    filter_type = filter_type[0]

                                ext = format_mapping.get(filter_type, "png")

                                # Save image
                                image_path = self.temp_dir / f"pdf_p{page_num}_img_{img_hash[:8]}.{ext}"
                                with open(image_path, 'wb') as img_file:
                                    img_file.write(data)

                                # Validate image
                                if await self._is_valid_image(image_path):
                                    image_files.append(image_path)
                                else:
                                    image_path.unlink()

                            except Exception as e:
                                self.logger.warning(f"Failed to extract PDF image on page {page_num}: {e}")
                                continue

            return image_files

        except Exception as e:
            self.logger.error(f"Error processing PDF images: {e}")
            return []

    async def _extract_docx_images(self, docx_path: Path) -> List[Path]:
        """
        Extract images from DOCX file.

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of paths to extracted images
        """
        image_files = []
        try:
            doc = Document(docx_path)

            # Process inline shapes
            for shape in doc.inline_shapes:
                try:
                    if shape.type == InlineShape.PICTURE:
                        image_path = await self._save_docx_image(shape)
                        if image_path:
                            image_files.append(image_path)
                except Exception as e:
                    self.logger.warning(f"Failed to process inline shape: {e}")

            # Process shapes from all paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if hasattr(run, '_element') and run._element.drawing_lst:
                        try:
                            for drawing in run._element.drawing_lst:
                                image_path = await self._save_docx_drawing(drawing)
                                if image_path:
                                    image_files.append(image_path)
                        except Exception as e:
                            self.logger.warning(f"Failed to process drawing: {e}")

            return image_files

        except Exception as e:
            self.logger.error(f"Error processing DOCX images: {e}")
            return []

    async def _save_docx_image(self, shape: InlineShape) -> Optional[Path]:
        """
        Save image from DOCX inline shape.

        Args:
            shape: DOCX inline shape object

        Returns:
            Path to saved image or None if failed
        """
        try:
            # Get image data
            if not hasattr(shape, '_inline'):
                return None

            if not hasattr(shape._inline, 'graphic'):
                return None

            image_data = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = shape.part.related_parts.get(image_data)

            if not isinstance(image_part, ImagePart):
                return None

            image_bytes = image_part.blob

            # Check uniqueness
            img_hash = hashlib.md5(image_bytes).hexdigest()
            if img_hash in self.image_hashes:
                return None
            self.image_hashes.add(img_hash)

            # Determine image format
            image = Image.open(io.BytesIO(image_bytes))
            ext = image.format.lower() if image.format else 'png'

            # Save image
            image_path = self.temp_dir / f"docx_img_{img_hash[:8]}.{ext}"
            with open(image_path, 'wb') as img_file:
                img_file.write(image_bytes)

            # Validate saved image
            if await self._is_valid_image(image_path):
                return image_path
            else:
                image_path.unlink()
                return None

        except Exception as e:
            self.logger.warning(f"Failed to save DOCX image: {e}")
            return None

    async def _save_docx_drawing(self, drawing) -> Optional[Path]:
        """
        Save image from DOCX drawing element.

        Args:
            drawing: DOCX drawing element

        Returns:
            Path to saved image or None if failed
        """
        try:
            # Extract image data from drawing
            image_data = None
            try:
                # Try to get image data from different possible locations
                if hasattr(drawing, 'graphic'):
                    if hasattr(drawing.graphic, 'graphicData'):
                        if hasattr(drawing.graphic.graphicData, 'pic'):
                            if hasattr(drawing.graphic.graphicData.pic, 'blipFill'):
                                if hasattr(drawing.graphic.graphicData.pic.blipFill, 'blip'):
                                    image_data = drawing.graphic.graphicData.pic.blipFill.blip.embed
            except AttributeError:
                return None

            if not image_data:
                return None

            # Get related part containing the image
            if not hasattr(drawing, 'part') or not hasattr(drawing.part, 'related_parts'):
                return None

            image_part = drawing.part.related_parts.get(image_data)
            if not isinstance(image_part, ImagePart):
                return None

            image_bytes = image_part.blob

            # Check uniqueness
            img_hash = hashlib.md5(image_bytes).hexdigest()
            if img_hash in self.image_hashes:
                return None
            self.image_hashes.add(img_hash)

            # Determine image format
            image = Image.open(io.BytesIO(image_bytes))
            ext = image.format.lower() if image.format else 'png'

            # Save image
            image_path = self.temp_dir / f"docx_drawing_{img_hash[:8]}.{ext}"
            with open(image_path, 'wb') as img_file:
                img_file.write(image_bytes)

            # Validate saved image
            if await self._is_valid_image(image_path):
                return image_path
            else:
                image_path.unlink()
                return None

        except Exception as e:
            self.logger.warning(f"Failed to save DOCX drawing: {e}")
            return None

    async def _is_valid_image(self, image_path: Path) -> bool:
        """
        Validate image file and check minimum size requirements.

        Args:
            image_path: Path to image file

        Returns:
            bool: True if image is valid, False otherwise
        """
        try:
            with Image.open(image_path) as img:
                # Check if image is too small
                if img.width < self.min_width or img.height < self.min_height:
                    return False

                # Check if image is empty or corrupted
                if img.getbbox() is None:
                    return False

                # Try to verify image data
                img.verify()

                return True

        except Exception as e:
            self.logger.warning(f"Invalid image {image_path}: {e}")
            return False

    def cleanup(self):
        """Remove temporary image files and directory."""
        try:
            # Remove all files in temp directory
            for file in self.temp_dir.glob('*'):
                try:
                    file.unlink()
                except Exception as e:
                    self.logger.warning(f"Failed to remove temp file {file}: {e}")

            # Remove temp directory
            try:
                self.temp_dir.rmdir()
            except Exception as e:
                self.logger.warning(f"Failed to remove temp directory {self.temp_dir}: {e}")

        except Exception as e:
            self.logger.error(f"Error during cleanup: {e}")

    async def get_image_data(self, image_path: Path) -> Optional[str]:
        """
        Get base64 encoded image data.

        Args:
            image_path: Path to image file

        Returns:
            Base64 encoded image data or None if failed
        """
        try:
            with open(image_path, 'rb') as img_file:
                image_data = img_file.read()
                return base64.b64encode(image_data).decode('utf-8')
        except Exception as e:
            self.logger.error(f"Failed to get image data for {image_path}: {e}")
            return None


class EnhancedConfigExtractor(ConfigExtractor):
    """Enhanced configuration extractor with image processing capabilities."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.image_extractor = ImageExtractor()

    async def process_document(self, file_path: Path, doc_type: str) -> Dict[str, Any]:
        """Process document including any embedded images."""
        try:
            # Extract text configurations
            text_configs = await super().process_document(file_path, doc_type)

            # Extract and process images
            image_configs = await self._process_document_images(file_path, doc_type)

            # Merge configurations
            final_config = self._merge_configs(text_configs, image_configs)

            return final_config

        finally:
            # Cleanup temporary image files
            self.image_extractor.cleanup()

    async def _process_document_images(self, file_path: Path, doc_type: str) -> Dict[str, Any]:
        """Extract and process images from document."""
        image_paths = await self.image_extractor.extract_images(file_path, doc_type)

        if not image_paths:
            return {}

        # Process images concurrently
        async with asyncio.TaskGroup() as group:
            image_tasks = [
                group.create_task(self._process_single_image(img_path))
                for img_path in image_paths
            ]

        # Collect and merge image results
        image_configs = {}
        for task in image_tasks:
            try:
                config = task.result()
                if config:
                    image_configs = self._merge_configs(image_configs, config)
            except Exception as e:
                self.logger.warning(f"Failed to process image configuration: {e}")

        return image_configs

    async def _process_single_image(self, image_path: Path) -> Dict[str, Any]:
        """Process a single image for configuration information."""
        try:
            # Convert image to base64
            with open(image_path, 'rb') as img_file:
                base64_image = base64.b64encode(img_file.read()).decode('utf-8')

            # Get configuration from image using GPT-4V
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": """Analyze this image for configuration information. 
                                     Look for:
                                     1. Database connection details
                                     2. Server configurations
                                     3. Network settings
                                     4. Environment variables
                                     5. System parameters
                                     Extract any values that match the configuration schema."""
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ]

            # Using the new OpenAI client format
            response = await self.openai_client.chat.completions.create(
                model=self.vision_model,
                messages=messages,
                max_tokens=1000
            )

            # Parse and validate the response
            extracted_config = self._parse_vision_response(
                response.choices[0].message.content
            )

            return extracted_config

        except Exception as e:
            self.logger.warning(f"Error processing image {image_path}: {e}")
            return {}

    def _merge_configs(self, config1: Dict[str, Any], config2: Dict[str, Any]) -> Dict[str, Any]:
        """Merge two configurations with conflict resolution."""
        merged = config1.copy()

        for key, value in config2.items():
            if key not in merged:
                merged[key] = value
            elif isinstance(value, dict) and isinstance(merged[key], dict):
                merged[key] = self._merge_configs(merged[key], value)
            else:
                # Keep the value with higher confidence or the more specific one
                merged[key] = self._resolve_config_conflict(merged[key], value)

        return merged

    def _resolve_config_conflict(self, value1: Any, value2: Any) -> Any:
        """Resolve conflicts between two configuration values."""
        # Implement conflict resolution logic here
        # For now, prefer non-None values and take the more specific value
        if value1 is None:
            return value2
        if value2 is None:
            return value1

        # Prefer more specific values (e.g., actual hostnames over "localhost")
        if isinstance(value1, str) and isinstance(value2, str):
            if value1.lower() in ['localhost', '127.0.0.1'] and value2.lower() not in ['localhost', '127.0.0.1']:
                return value2
            if value2.lower() in ['localhost', '127.0.0.1'] and value1.lower() not in ['localhost', '127.0.0.1']:
                return value1

        # Default to first value
        return value1


# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'docx',
    '.mp4': 'video',
    '.avi': 'video',
    '.mov': 'video',
    '.mkv': 'video'
}


class DocumentProcessor:
    def __init__(self):
        self.console = Console()
        self.config = config.load_config()
        self.config_extractor = self._initialize_extractor()
        self.logger = self._setup_logger()

    def _initialize_extractor(self) -> ConfigExtractor:
        """Initialize ConfigExtractor with schema from config."""
        try:
            schema_path = self.config.processing.schema_path
            if not os.path.exists(schema_path):
                self.console.print(f"[yellow]Warning: Schema file not found at {schema_path}[/yellow]")
                return ConfigExtractor(model_name=self.config.openai.gpt4_model)

            return ConfigExtractor(
                model_name=self.config.openai.gpt4_model,
                schema_path=schema_path
            )
        except Exception as e:
            self.console.print(f"[red]Error loading schema: {e}[/red]")
            return ConfigExtractor(model_name=self.config.openai.gpt4_model)

    def _setup_logger(self) -> logging.Logger:
        """Setup logging configuration."""
        logger = logging.getLogger("DocumentProcessor")
        logger.setLevel(logging.INFO)

        # Create logs directory if it doesn't exist
        os.makedirs(self.config.log_dir, exist_ok=True)

        # File handler
        log_file = os.path.join(self.config.log_dir, f"extraction_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)

        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)

        # Formatter
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)

        logger.addHandler(file_handler)
        logger.addHandler(console_handler)

        return logger

    async def process_directory(self, directory_path: str | Path) -> Dict[str, Any]:
        """Process all supported files in the directory concurrently."""
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # Create output and temp directories if they don't exist
        os.makedirs(self.config.output_dir, exist_ok=True)
        os.makedirs(self.config.temp_dir, exist_ok=True)

        # Get list of supported files
        files_to_process = []
        for ext, doc_type in SUPPORTED_EXTENSIONS.items():
            files_to_process.extend(
                (file, doc_type) for file in directory_path.glob(f"**/*{ext}")
            )

        if not files_to_process:
            self.logger.warning("No supported files found in directory")
            return {}

        # Process files with progress tracking
        results = {}
        with Progress() as progress:
            total_progress = progress.add_task("[cyan]Processing files...", total=len(files_to_process))
            
            # Process files in batches
            batch_size = self.config.processing.concurrent_files
            for i in range(0, len(files_to_process), batch_size):
                batch = files_to_process[i:i + batch_size]
                
                # Create tasks for batch processing
                tasks = []
                task_ids = []
                for file_path, doc_type in batch:
                    task_id = progress.add_task(f"[green]Processing {file_path.name}", start=False)
                    task_ids.append(task_id)
                    tasks.append(self.process_file(file_path, doc_type, progress, task_id))

                # Process batch concurrently
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                
                # Update results
                for (file_path, doc_type), result in zip(batch, batch_results):
                    if isinstance(result, Exception):
                        self.logger.error(f"Error processing {file_path}: {str(result)}")
                        results[str(file_path)] = None
                    else:
                        processed_path, config_data = result
                        results[str(processed_path)] = config_data
                    progress.update(total_progress, advance=1)

                # Clean up task bars
                for task_id in task_ids:
                    progress.remove_task(task_id)

        return results

    async def process_file(
            self,
            file_path: Path,
            doc_type: str,
            progress: Progress,
            task_id: int
    ) -> tuple[Path, Dict[str, Any]]:
        """Process a single file with progress tracking."""
        try:
            progress.update(task_id, visible=True)

            # Extract text from document
            text = await self.config_extractor._read_document(file_path, doc_type)
            progress.update(task_id, description=f"[green]Text extracted {file_path.name}", completed=20)

            # First pass analysis
            first_pass_results = await self.config_extractor._first_pass_analysis(text)
            progress.update(task_id, description=f"[green]First pass completed {file_path.name}", completed=40)

            # Structure analysis
            structure_analysis = await self.config_extractor._analyze_structure(text)
            progress.update(task_id, description=f"[green]Structure analyzed {file_path.name}", completed=60)

            # Find related contexts
            related_contexts = await self.config_extractor._find_related_contexts(text)
            progress.update(task_id, description=f"[green]Contexts found {file_path.name}", completed=80)

            # Second pass extraction
            config_result = await self.config_extractor._second_pass_extraction(
                text,
                first_pass_results,
                structure_analysis,
                related_contexts
            )
            config_result = await self.config_extractor.process_document(file_path, doc_type)

            progress.update(task_id, description=f"[green]Completed {file_path.name}", completed=100)
            return file_path, config_result

        except Exception as e:
            #log error detail
            progress.update(task_id, description=f"[red]Failed {file_path.name}")
            raise RuntimeError(f"Error processing {file_path}: {str(e)}")

        finally:
            # Hide the task once it's done
            progress.update(task_id, visible=False)


async def main():
    """Main entry point for the document processing script."""
    import argparse
    from pathlib import Path
    import json
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TimeElapsedColumn

    # Set up argument parser with better help messages
    parser = argparse.ArgumentParser(
        description="Configuration Extraction from Multiple Documents",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument(
        'directory',
        type=str,
        help="Directory containing documents to process"
    )
    parser.add_argument(
        '--output', '-o',
        type=str,
        required=True,
        help="Output file path for results (JSON format)"
    )
    parser.add_argument(
        '--schema', '-s',
        type=str,
        default="./schema/config_schema.json",
        help="Path to schema file"
    )
    parser.add_argument(
        '--concurrent-files', '-c',
        type=int,
        default=5,
        help="Number of files to process concurrently"
    )

    try:
        # Parse arguments
        args = parser.parse_args()
        console = Console()

        # Convert directory path to Path object and validate
        input_dir = Path(args.directory)
        output_file = Path(args.output)

        # Ensure input directory exists
        if not input_dir.exists() or not input_dir.is_dir():
            console.print(f"\n[bold red]Error: Directory not found: {input_dir}[/bold red]")
            return 1

        # Ensure output directory exists, create if not
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # Load configuration
        config_instance = config.load_config()

        # Update config with command line arguments
        config_instance.processing.schema_path = args.schema
        config_instance.processing.concurrent_files = args.concurrent_files

        # Process start time
        start_time = time.time()

        console.print("\n[bold cyan]Starting document processing...[/bold cyan]")

        # Initialize processor with updated configuration
        processor = DocumentProcessor()

        # Get list of all supported files
        files_to_process = []
        for ext, doc_type in SUPPORTED_EXTENSIONS.items():
            files_to_process.extend(
                (file, doc_type) for file in input_dir.glob(f"**/*{ext}")
            )

        if not files_to_process:
            console.print("[yellow]No supported files found in directory[/yellow]")
            return 0

        # Process files with progress tracking
        results = {}
        with Progress(
                SpinnerColumn(),
                *Progress.get_default_columns(),
                TimeElapsedColumn(),
                console=console
        ) as progress:
            # Create overall progress bar
            total_progress = progress.add_task(
                "[cyan]Processing files...",
                total=len(files_to_process)
            )

            # Process files in batches
            batch_size = config_instance.processing.concurrent_files
            for i in range(0, len(files_to_process), batch_size):
                batch = files_to_process[i:i + batch_size]

                # Create tasks for each file in batch
                tasks = []
                task_ids = []
                for file_path, doc_type in batch:
                    task_id = progress.add_task(
                        f"[green]Processing {file_path.name}",
                        start=False
                    )
                    task_ids.append(task_id)
                    tasks.append(processor.process_file(file_path, doc_type, progress, task_id))

                # Process batch concurrently
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)

                # Update results and progress
                for (file_path, doc_type), result in zip(batch, batch_results):
                    if isinstance(result, Exception):
                        console.print(f"[red]Error processing {file_path}: {str(result)}[/red]")
                        results[str(file_path)] = None
                    elif isinstance(result, tuple):
                        processed_path, config_data = result
                        results[str(processed_path)] = config_data
                    progress.update(total_progress, advance=1)

                # Remove completed task bars
                for task_id in task_ids:
                    progress.remove_task(task_id)

        # Write results to output file
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            console.print(f"\n[green]Results written to {output_file}[/green]")
        except Exception as e:
            console.print(f"\n[bold red]Error writing output file: {str(e)}[/bold red]")
            return 1

        # Print summary
        processing_time = time.time() - start_time
        total_files = len(files_to_process)
        successful_files = sum(1 for result in results.values() if result is not None)

        console.print("\n[bold cyan]Processing Summary:[/bold cyan]")
        console.print(f"Total files processed: {total_files}")
        console.print(f"Successfully processed: {successful_files}")
        console.print(f"Failed: {total_files - successful_files}")
        console.print(f"Total processing time: {processing_time:.2f} seconds")
        if total_files > 0:
            console.print(f"Average time per file: {processing_time / total_files:.2f} seconds")

        return 0

    except KeyboardInterrupt:
        console.print("\n[yellow]Processing interrupted by user[/yellow]")
        return 1
    except Exception as e:
        console.print(f"\n[bold red]Fatal error: {str(e)}[/bold red]")
        return 1


if __name__ == "__main__":
    try:
        if sys.platform.startswith('win'):
            # Windows specific event loop policy
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
        exit_code = asyncio.run(main())
        sys.exit(exit_code)
    except KeyboardInterrupt:
        Console().print("\n[yellow]Processing interrupted by user[/yellow]")
        sys.exit(1)
    except Exception as e:
        Console().print(f"\n[bold red]Fatal error: {str(e)}[/bold red]")
        sys.exit(1)
